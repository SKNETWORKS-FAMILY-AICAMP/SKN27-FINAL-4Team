from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\Playdata\Downloads\[데이터 수집 및 저장] 데이터베이스_저장소 설계 문서_27기_0팀.docx")
OUT_DIR = ROOT / "docs" / "한재웅"
OUTPUT = OUT_DIR / "[데이터 수집 및 저장] 데이터베이스_저장소 설계 문서_27기_4팀_작성본.docx"
ERD = OUT_DIR / "postgresql_erd_27_4team.png"

FONT = "맑은 고딕"
FONT_PATH = Path(r"C:\Windows\Fonts\malgun.ttf")
BOLD_FONT_PATH = Path(r"C:\Windows\Fonts\malgunbd.ttf")


def set_run_font(run, size=9, bold=False, color=None):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), FONT)


def set_para_text(paragraph, text, size=10, bold=False, align=None):
    paragraph.clear()
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if align is not None:
        paragraph.alignment = align


def set_cell_text(cell, text, size=8.5, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, shade=None):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(str(text).split("\n")):
        if idx:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=size, bold=bold)
    if shade:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            tc_pr.append(shd)
        shd.set(qn("w:fill"), shade)


def clear_table_keep_header(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def fill_table(table, rows, header=True):
    clear_table_keep_header(table)
    needed_cols = len(table.columns)
    for c_idx, value in enumerate(rows[0][:needed_cols]):
        set_cell_text(table.rows[0].cells[c_idx], value, size=8.5, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, shade="D9EAF7")
    for row in rows[1:]:
        cells = table.add_row().cells
        for c_idx in range(needed_cols):
            text = row[c_idx] if c_idx < len(row) else ""
            center = c_idx in {0, 2, 3, 4}
            set_cell_text(cells[c_idx], text, size=8, align=WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT)


def insert_paragraph_after(anchor, text="", style=None, size=10, bold=False):
    new_p = OxmlElement("w:p")
    anchor._element.addnext(new_p)
    paragraph = anchor._parent.add_paragraph()
    paragraph._p = new_p
    paragraph._element = new_p
    if style:
        paragraph.style = style
    if text:
        set_para_text(paragraph, text, size=size, bold=bold)
    return paragraph


def clone_table_after(doc, anchor_table, title, rows):
    title_p = doc.add_paragraph()
    set_para_text(title_p, title, size=10, bold=True)
    new_tbl = deepcopy(anchor_table._tbl)
    anchor_table._tbl.addnext(title_p._p)
    title_p._p.addnext(new_tbl)
    table = Table(new_tbl, anchor_table._parent)
    fill_table(table, rows)
    return table


def add_doc_table_after(doc, anchor_element, title, rows):
    title_p = doc.add_paragraph()
    set_para_text(title_p, title, size=10, bold=True)
    tbl = doc.add_table(rows=1, cols=len(rows[0]))
    tbl.style = "Table Grid"
    fill_table(tbl, rows)
    anchor_element.addnext(title_p._p)
    title_p._p.addnext(tbl._tbl)
    return tbl


def draw_box(draw, xy, title, lines, fill, outline="#476582"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2)
    title_font = ImageFont.truetype(str(BOLD_FONT_PATH), 20)
    body_font = ImageFont.truetype(str(FONT_PATH), 15)
    draw.text((x1 + 14, y1 + 12), title, font=title_font, fill="#18324A")
    y = y1 + 44
    for line in lines:
        draw.text((x1 + 14, y), line, font=body_font, fill="#263238")
        y += 22


def line(draw, start, end, label=None):
    draw.line([start, end], fill="#4E6E8E", width=3)
    ex, ey = end
    sx, sy = start
    # Simple arrow head
    if ex >= sx:
        points = [(ex, ey), (ex - 10, ey - 6), (ex - 10, ey + 6)]
    else:
        points = [(ex, ey), (ex + 10, ey - 6), (ex + 10, ey + 6)]
    draw.polygon(points, fill="#4E6E8E")
    if label:
        font = ImageFont.truetype(str(FONT_PATH), 14)
        lx = (sx + ex) // 2 - 18
        ly = (sy + ey) // 2 - 22
        draw.rounded_rectangle((lx - 4, ly - 2, lx + 58, ly + 20), radius=4, fill="#FFFFFF", outline="#C7D2DD")
        draw.text((lx, ly), label, font=font, fill="#345")


def draw_entity(draw, xy, title, fields, fill="#FFFFFF"):
    x1, y1, x2, y2 = xy
    header_h = 38
    draw.rounded_rectangle(xy, radius=8, fill=fill, outline="#2F567C", width=2)
    draw.rounded_rectangle((x1, y1, x2, y1 + header_h), radius=8, fill="#D9EAF7", outline="#2F567C", width=2)
    # Flatten the bottom of the header so it reads like a table header.
    draw.rectangle((x1, y1 + header_h - 8, x2, y1 + header_h), fill="#D9EAF7", outline="#2F567C")
    title_font = ImageFont.truetype(str(BOLD_FONT_PATH), 17)
    field_font = ImageFont.truetype(str(FONT_PATH), 13)
    draw.text((x1 + 10, y1 + 8), title, font=title_font, fill="#17324D")
    y = y1 + header_h + 8
    for field in fields:
        draw.text((x1 + 10, y), field, font=field_font, fill="#243746")
        y += 20


def draw_relationship(draw, start, end, start_card, end_card, label=None):
    sx, sy = start
    ex, ey = end
    draw.line([start, end], fill="#456C8E", width=3)
    font = ImageFont.truetype(str(BOLD_FONT_PATH), 13)
    label_font = ImageFont.truetype(str(FONT_PATH), 12)
    draw.rounded_rectangle((sx - 20, sy - 12, sx + 20, sy + 10), radius=4, fill="#FFFFFF", outline="#BAC8D6")
    draw.text((sx - 13, sy - 10), start_card, font=font, fill="#17324D")
    draw.rounded_rectangle((ex - 20, ey - 12, ex + 22, ey + 10), radius=4, fill="#FFFFFF", outline="#BAC8D6")
    draw.text((ex - 14, ey - 10), end_card, font=font, fill="#17324D")
    if label:
        mx, my = (sx + ex) // 2, (sy + ey) // 2
        tw = max(44, len(label) * 9)
        draw.rounded_rectangle((mx - tw // 2, my - 13, mx + tw // 2, my + 10), radius=4, fill="#FFFFFF", outline="#D7E0EA")
        draw.text((mx - tw // 2 + 6, my - 10), label, font=label_font, fill="#345")


def draw_relationship_path(draw, points, start_card, end_card, label=None):
    for a, b in zip(points, points[1:]):
        draw.line([a, b], fill="#456C8E", width=3)
    font = ImageFont.truetype(str(BOLD_FONT_PATH), 13)
    label_font = ImageFont.truetype(str(FONT_PATH), 12)
    sx, sy = points[0]
    ex, ey = points[-1]
    draw.rounded_rectangle((sx - 20, sy - 12, sx + 20, sy + 10), radius=4, fill="#FFFFFF", outline="#BAC8D6")
    draw.text((sx - 13, sy - 10), start_card, font=font, fill="#17324D")
    draw.rounded_rectangle((ex - 20, ey - 12, ex + 22, ey + 10), radius=4, fill="#FFFFFF", outline="#BAC8D6")
    draw.text((ex - 14, ey - 10), end_card, font=font, fill="#17324D")
    if label:
        mid = points[len(points) // 2]
        mx, my = mid
        tw = max(44, len(label) * 9)
        draw.rounded_rectangle((mx - tw // 2, my - 13, mx + tw // 2, my + 10), radius=4, fill="#FFFFFF", outline="#D7E0EA")
        draw.text((mx - tw // 2 + 6, my - 10), label, font=label_font, fill="#345")


def draw_bus_relation(draw, bus_x, user_y, child_point, label, end_card="N"):
    label_font = ImageFont.truetype(str(FONT_PATH), 12)
    card_font = ImageFont.truetype(str(BOLD_FONT_PATH), 13)
    cx, cy = child_point
    draw.line([(430, user_y), (bus_x, user_y), (bus_x, cy), (cx, cy)], fill="#456C8E", width=3)
    tw = max(44, len(label) * 8)
    lx = (bus_x + cx) // 2
    draw.rounded_rectangle((lx - tw // 2, cy - 13, lx + tw // 2, cy + 10), radius=4, fill="#FFFFFF", outline="#D7E0EA")
    draw.text((lx - tw // 2 + 6, cy - 10), label, font=label_font, fill="#345")
    draw.rounded_rectangle((cx - 20, cy - 12, cx + 22, cy + 10), radius=4, fill="#FFFFFF", outline="#BAC8D6")
    draw.text((cx - 14, cy - 10), end_card, font=card_font, fill="#17324D")


def make_erd():
    img = Image.new("RGB", (1800, 1250), "#F7FAFC")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(str(BOLD_FONT_PATH), 32)
    sub_font = ImageFont.truetype(str(FONT_PATH), 18)
    draw.text((50, 34), "Mind Wellness PostgreSQL ERD", font=title_font, fill="#18324A")
    draw.text((52, 76), "Django 모델 및 ETL 스키마 기준 엔터티 관계도", font=sub_font, fill="#52616E")

    entities = {
        "users": ((80, 520, 430, 680), "USERS", ["PK id", "UQ email", "nickname", "character", "onboarding_done", "created_at"]),
        "oauth": ((700, 90, 1090, 240), "OAUTH_ACCOUNTS", ["PK id", "FK user_id", "provider", "provider_user_id", "raw_profile JSON"]),
        "profile": ((700, 280, 1090, 430), "USER_PROFILES", ["PK id", "FK user_id", "birth_date", "gender, age, job", "hobbies/interests JSON"]),
        "pref_kw": ((700, 470, 1090, 620), "USER_PREFERENCE_KEYWORDS", ["PK id", "FK user_id", "keyword_type", "label", "source"]),
        "session": ((700, 660, 1090, 825), "CHAT_SESSIONS", ["PK id", "FK user_id", "character", "is_secret", "selected_emotion"]),
        "memory": ((700, 865, 1090, 980), "USER_MEMORY", ["PK/FK user_id", "summary_text", "updated_at"]),
        "tarot": ((700, 1020, 1090, 1165), "DAILY_TAROT_FORTUNES", ["PK id", "FK user_id", "target_date", "fortune_type", "card_number"]),
        "mbti_q": ((1240, 90, 1660, 260), "MBTI_QUESTION_RESPONSES", ["PK id", "user_id", "target_axis", "period_key", "answer_text"]),
        "mbti_r": ((1240, 310, 1660, 470), "MBTI_MONTHLY_RESULTS", ["PK id", "user_id", "period_key", "estimated_mbti_type", "status"]),
        "scale_est": ((1240, 520, 1660, 690), "USER_SCALE_ESTIMATES", ["PK id", "FK user_id", "FK scale_id", "estimated_score", "target_date"]),
        "scale": ((1240, 730, 1660, 850), "CLINICAL_SCALES", ["PK scale_id", "scale_name_ko", "domain", "time_frame"]),
        "message": ((1240, 890, 1660, 1060), "CHAT_MESSAGES", ["PK id", "FK session_id", "role", "content", "emotion_label", "created_at"]),
        "taste": ((1240, 1100, 1660, 1235), "TASTE_ANALYSIS", ["CONVERSATION_LOGS", "PREFERENCE_EVIDENCE", "PREFERENCE_SUMMARIES", "user_id / message_id"]),
    }

    for key, (xy, title, fields) in entities.items():
        draw_entity(draw, xy, title, fields)

    bus_x = 560
    draw.line([(430, 600), (bus_x, 600)], fill="#456C8E", width=3)
    draw.line([(bus_x, 160), (bus_x, 1168)], fill="#456C8E", width=3)
    card_font = ImageFont.truetype(str(BOLD_FONT_PATH), 13)
    draw.rounded_rectangle((425, 588, 466, 612), radius=4, fill="#FFFFFF", outline="#BAC8D6")
    draw.text((439, 590), "1", font=card_font, fill="#17324D")

    draw_bus_relation(draw, bus_x, 600, (700, 165), "auth")
    draw_bus_relation(draw, bus_x, 600, (700, 355), "profile", "1")
    draw_bus_relation(draw, bus_x, 600, (700, 545), "keywords")
    draw_bus_relation(draw, bus_x, 600, (700, 742), "sessions")
    draw_bus_relation(draw, bus_x, 600, (700, 922), "memory", "1")
    draw_bus_relation(draw, bus_x, 600, (700, 1092), "tarot")
    draw_bus_relation(draw, bus_x, 600, (1240, 175), "MBTI Q")
    draw_bus_relation(draw, bus_x, 600, (1240, 390), "monthly")
    draw_bus_relation(draw, bus_x, 600, (1240, 605), "scale result")
    draw_bus_relation(draw, bus_x, 600, (1240, 1168), "taste")

    draw_relationship(draw, (1090, 742), (1240, 975), "1", "N", "messages")
    draw_relationship(draw, (1450, 730), (1450, 690), "1", "N", "scale")

    img.save(ERD)


ENTITY_ROWS = [
    ["구분", "테이블", "제약 대상", "제약 유형", "주요 속성"],
    ["계정", "users", "사용자", "PK, UNIQUE", "email, nickname, character, onboarding_done"],
    ["계정", "oauth_accounts", "소셜 로그인", "FK, UNIQUE", "user_id, provider, provider_user_id, raw_profile"],
    ["프로필", "user_profiles", "온보딩 프로필", "1:1 FK", "birth_date, gender, age, hobbies, interests"],
    ["프로필", "user_preference_keywords", "취향 키워드", "FK, UNIQUE", "keyword_type, label, source"],
    ["대화", "chat_sessions", "대화 세션", "FK", "character, is_secret, selected_emotion, mbti_pending"],
    ["대화", "chat_messages", "대화 메시지", "FK", "role, content, emotion_label, created_at"],
    ["대화", "user_memory", "장기 요약", "PK/FK", "summary_text, updated_at"],
    ["MBTI", "mbti_question_responses", "월간 문답 원천", "INDEX", "target_axis, period_key, answered_at"],
    ["MBTI", "mbti_response_scores", "문답 채점 결과", "1:1 FK", "axis, score, direction, coding_status"],
    ["MBTI", "mbti_monthly_results", "월간 최종 결과", "UNIQUE", "user_id, period_key, estimated_mbti_type"],
    ["웰니스", "clinical_scales", "척도 마스터", "PK", "scale_id, scale_name_ko, domain, time_frame"],
    ["웰니스", "user_scale_estimates", "사용자 척도 추정", "FK", "user_id, scale_id, estimated_score, target_date"],
    ["운세", "daily_tarot_fortunes", "일일 타로 결과", "FK, UNIQUE", "user_id, target_date, card_number, source"],
    ["취향", "preference_evidence", "취향 근거", "FK", "message_id, normalized_keyword, evidence_text"],
]

REL_ROWS = [
    ["관계명", "주 엔터티", "종 엔터티", "관계"],
    ["사용자 - 소셜 계정", "users", "oauth_accounts", "1:N"],
    ["사용자 - 프로필", "users", "user_profiles", "1:1"],
    ["사용자 - 온보딩 키워드", "users", "user_preference_keywords", "1:N"],
    ["사용자 - 대화 세션", "users", "chat_sessions", "1:N"],
    ["대화 세션 - 메시지", "chat_sessions", "chat_messages", "1:N"],
    ["사용자 - 장기 요약", "users", "user_memory", "1:1"],
    ["MBTI 문답 - 채점", "mbti_question_responses", "mbti_response_scores", "1:1"],
    ["월간 MBTI 결과 - 축별 결과", "mbti_monthly_results", "mbti_monthly_axis_results", "1:N"],
    ["월간 MBTI 결과 - 리포트", "mbti_monthly_results", "mbti_monthly_reports", "1:1"],
    ["척도 - 문항/선택지", "clinical_scales", "scale_questions / scale_options", "1:N"],
    ["사용자 - 척도 추정", "users", "user_scale_estimates", "1:N"],
    ["사용자 - 일일 운세", "users", "daily_fortunes / daily_tarot_fortunes", "1:N"],
    ["대화 로그 - 취향 근거", "conversation_logs", "preference_evidence", "1:N"],
]

USER_TABLE = [
    ["컬럼명", "타입", "PK", "FK", "Not Null", "제약 조건 설명"],
    ["id", "BIGSERIAL", "Y", "", "Y", "Django 기본 사용자 고유 식별자"],
    ["email", "VARCHAR", "", "", "Y", "UNIQUE, 로그인 ID 및 OAuth 매칭 기준"],
    ["password", "VARCHAR", "", "", "Y", "Django 해시 비밀번호"],
    ["nickname", "VARCHAR(30)", "", "", "Y", "서비스 표시 이름"],
    ["character", "VARCHAR(10)", "", "", "", "사용자 선택 캐릭터 코드"],
    ["onboarding_done", "BOOLEAN", "", "", "Y", "온보딩 완료 여부"],
    ["is_active / is_staff", "BOOLEAN", "", "", "Y", "계정 활성/관리자 권한 플래그"],
    ["created_at", "TIMESTAMPTZ", "", "", "Y", "가입 일시(auto_now_add)"],
]

TABLE_DEFS = {
    "테이블 : chat_sessions": [
        ["컬럼명", "타입", "PK", "FK", "Not Null", "제약 조건 설명"],
        ["id", "BIGSERIAL", "Y", "", "Y", "대화 세션 식별자"],
        ["user_id", "BIGINT", "", "Y", "", "users.id 참조, 비로그인/익명 흐름 허용"],
        ["character", "VARCHAR(10)", "", "", "Y", "선택 캐릭터 코드"],
        ["is_secret", "BOOLEAN", "", "", "Y", "시크릿챗 여부, 저장 정책 분기"],
        ["selected_emotion", "VARCHAR(10)", "", "", "", "초기 감정 선택값"],
        ["mbti_pending", "BOOLEAN", "", "", "Y", "MBTI 후속 질문 대기 상태"],
        ["created_at", "TIMESTAMPTZ", "", "", "Y", "세션 생성 일시"],
    ],
    "테이블 : chat_messages": [
        ["컬럼명", "타입", "PK", "FK", "Not Null", "제약 조건 설명"],
        ["id", "BIGSERIAL", "Y", "", "Y", "메시지 식별자"],
        ["session_id", "BIGINT", "", "Y", "Y", "chat_sessions.id 참조, CASCADE 삭제"],
        ["role", "VARCHAR(10)", "", "", "Y", "user / assistant"],
        ["content", "TEXT", "", "", "Y", "대화 원문"],
        ["emotion_label", "VARCHAR(20)", "", "", "", "joy/sadness/anger/normal 분류값"],
        ["created_at", "TIMESTAMPTZ", "", "", "Y", "메시지 생성 일시"],
    ],
    "테이블 : mbti_monthly_results": [
        ["컬럼명", "타입", "PK", "FK", "Not Null", "제약 조건 설명"],
        ["id", "BIGSERIAL", "Y", "", "Y", "월간 분석 결과 식별자"],
        ["user_id", "BIGINT", "", "", "Y", "사용자 식별자, 기간 조회 인덱스"],
        ["period_key", "VARCHAR(7)", "", "", "Y", "YYYY-MM 분석 월"],
        ["estimated_mbti_type", "VARCHAR(4)", "", "", "", "월간 추정 MBTI 유형"],
        ["changed_axes_json", "JSONB", "", "", "Y", "전월 대비 변화 축 목록"],
        ["status", "VARCHAR(32)", "", "", "Y", "complete/failed 등 처리 상태"],
        ["analyzed_at", "TIMESTAMPTZ", "", "", "", "분석 완료 시각"],
    ],
    "테이블 : clinical_scales": [
        ["컬럼명", "타입", "PK", "FK", "Not Null", "제약 조건 설명"],
        ["scale_id", "VARCHAR(50)", "Y", "", "Y", "PHQ-9, GAD-7 등 척도 코드"],
        ["scale_name_ko", "VARCHAR(100)", "", "", "Y", "한글 척도명"],
        ["domain", "VARCHAR(50)", "", "", "Y", "mental / physical 등 도메인"],
        ["time_frame", "VARCHAR(100)", "", "", "", "문항 기준 기간"],
        ["estimated_minutes", "INTEGER", "", "", "Y", "예상 소요 시간"],
    ],
    "테이블 : daily_tarot_fortunes": [
        ["컬럼명", "타입", "PK", "FK", "Not Null", "제약 조건 설명"],
        ["id", "BIGSERIAL", "Y", "", "Y", "일일 타로 결과 식별자"],
        ["user_id", "BIGINT", "", "Y", "Y", "users.id 참조"],
        ["target_date", "DATE", "", "", "Y", "운세 대상 일자"],
        ["fortune_type", "VARCHAR(30)", "", "", "Y", "daily_major 등 운세 유형"],
        ["card_number", "INTEGER", "", "", "Y", "선택 카드 번호"],
        ["message", "TEXT", "", "", "", "사용자에게 제공되는 해석 문구"],
        ["source", "VARCHAR(30)", "", "", "Y", "rule/llm/hybrid 생성 방식"],
    ],
}

CONSTRAINT_ROWS = [
    ["적용 레벨", "제약 유형", "대상", "설명 및 관리 방안"],
    ["DB 레벨", "PK", "모든 Django 모델 테이블", "BigAutoField 또는 명시 PK로 레코드 식별성 보장"],
    ["DB 레벨", "FK", "oauth_accounts.user_id", "users.id 참조, ON DELETE CASCADE"],
    ["DB 레벨", "FK", "chat_sessions.user_id", "users.id 참조, 로그인 사용자의 대화 소유권 보장"],
    ["DB 레벨", "FK", "chat_messages.session_id", "chat_sessions.id 참조, 세션 삭제 시 메시지 동시 삭제"],
    ["DB 레벨", "FK", "user_memory.user_id", "users.id와 1:1, 사용자별 장기 요약 중복 방지"],
    ["DB 레벨", "UNIQUE", "users.email", "동일 이메일 중복 가입 방지"],
    ["DB 레벨", "UNIQUE", "oauth_accounts(provider, provider_user_id)", "소셜 계정 중복 연결 방지"],
    ["DB 레벨", "UNIQUE", "mbti_monthly_results(user_id, period_key)", "사용자/월별 MBTI 최종 결과 단일화"],
    ["DB 레벨", "UNIQUE", "daily_tarot_fortunes(user_id, target_date, fortune_type)", "동일 날짜/유형 운세 재생성 중복 방지"],
    ["DB 레벨", "INDEX", "user_id, period_key, target_axis", "월간 분석 및 마이페이지 조회 성능 최적화"],
    ["애플리케이션 레벨", "Secret Chat", "chat_sessions.is_secret", "시크릿 모드는 영구 저장 대상에서 제외하고 서버 메모리 기반 처리"],
    ["운영 레벨", "Migration", "Django migrations", "스키마 변경은 마이그레이션 파일로 이력 관리"],
]


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_erd()

    doc = Document(SOURCE)
    set_cell_text(doc.tables[0].rows[0].cells[0], "SK 네트웍스 Family AI 27기 : 4팀 / 데이터 베이스/저장소 설계 문서", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    meta = doc.tables[1]
    set_cell_text(meta.rows[0].cells[1], "데이터 수집 및 저장", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(meta.rows[1].cells[1], "2026. 07. 03.", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(meta.rows[2].cells[1], "C:\\dev\\project\\SKN27-FINAL-4Team", size=8)
    set_cell_text(meta.rows[3].cells[1], "27기 4팀", size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(doc.tables[2].rows[0].cells[0], "관계형 데이터베이스 (PostgreSQL 15 / Django ORM 기반)", size=11, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    set_para_text(doc.paragraphs[6], "모델링 방법", size=12, bold=True)
    set_para_text(doc.paragraphs[7], "모델링 방법 : Django 모델·마이그레이션을 기준으로 역공학하고, ETL SQL 스크립트와 화면/기획 문서의 저장 요구사항을 교차 검증하였다.", size=10)
    set_para_text(doc.paragraphs[8], "정규화 수준 : 사용자, 대화, MBTI 분석, 척도, 타로/운세, 취향 분석 도메인을 분리하여 3NF 수준을 기본으로 하되 JSONField는 화면 표시용 스냅샷·근거 묶음처럼 구조 변동이 큰 데이터에 한정하였다.", size=10)
    set_para_text(doc.paragraphs[9], "도구 : PostgreSQL 15, Django ORM, DBeaver ERD, diagrams.net(draw.io), Docker Compose, 프로젝트 migrations/models/ETL 스크립트", size=10)

    fill_table(doc.tables[3], ENTITY_ROWS)
    fill_table(doc.tables[4], REL_ROWS)
    fill_table(doc.tables[5], USER_TABLE)
    fill_table(doc.tables[6], CONSTRAINT_ROWS)
    fill_table(doc.tables[7], [
        ["변경일", "변경자", "변경내용", "영향 받는 항목", "비고"],
        ["2026.07.03", "27기 4팀", "PostgreSQL 단일 DB 기준으로 스키마 재정리", "전체", "docker-compose.yml v6.0 반영"],
        ["2026.07.03", "27기 4팀", "Django 모델 기반 테이블/제약조건/관계 갱신", "users, chat, mbti, wellness, tarot, taste", "models.py 및 migrations 참조"],
        ["2026.07.03", "27기 4팀", "ERD 삽입 및 핵심 도메인 관계 시각화", "2.3 ERD", "사용자 중심 핵심 관계도"],
    ])

    # Insert ERD image immediately after "2.3 ERD" with a report-style caption below.
    erd_anchor = doc.paragraphs[15]
    pic_p = insert_paragraph_after(erd_anchor, "", size=9)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(ERD), width=Inches(6.35))
    caption_p = insert_paragraph_after(pic_p, "그림 1. PostgreSQL 엔터티 관계도(ERD)", size=8, bold=True)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add additional table definitions after the original users table.
    anchor_tbl = doc.tables[5]
    current = anchor_tbl
    for title, rows in TABLE_DEFS.items():
        current = clone_table_after(doc, current, title, rows)

    doc.save(OUTPUT)
    print(OUTPUT)
    print(ERD)


if __name__ == "__main__":
    main()
